import streamlit as st
from docx import Document
from PIL import Image
import io
import pandas as pd
from docx.shared import Inches

st. set_page_config(layout="wide")

def extract_images(document):
    images = []
    for rel in document.part.rels.values():
        if "image" in rel.reltype:
            # Extract the image blob
            image_blob = rel.target_part.blob
            image_stream = io.BytesIO(image_blob)
            image = Image.open(image_stream)
            images.append(image)
    return images

def load_document(filepath):
    document = Document(filepath)

    content_by_header = {}
    current_header = None

    for para in document.paragraphs:
        if para.style.name.startswith('Heading'):
            current_header = para.text
            content_by_header[current_header] = {"paragraphs": [], "tables": [], "figures": []}
        elif current_header is not None:
            content_by_header[current_header]["paragraphs"].append(para.text)

    for rel in document.part.rels.values():
        if "image" in rel.reltype:
            # Extract the image blob
            image_blob = rel.target_part.blob
            image_stream = io.BytesIO(image_blob)
            image = Image.open(image_stream)
            if current_header:  # Add image under the current header
                content_by_header[current_header]["figures"].append(image)

    for table in document.tables:
        df_table = [['' for cell in row.cells] for row in table.rows]  # Create an empty 2D list for table data
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                df_table[i][j] = cell.text.strip()
        df = pd.DataFrame(df_table)
        if current_header:  # Append table under the corresponding header
            content_by_header[current_header]["tables"].append(df)

    return content_by_header

# def display_side_by_side(content1, content2):
#     col1, col2 = st.columns(2)
#     all_headers = set(content1.keys()) | set(content2.keys())
#     for header in sorted(all_headers):
#         with col1:
#             st.subheader(header)
#             for para in content1.get(header, {}).get("paragraphs", []):
#                 st.write(para)
#             for table in content1.get(header, {}).get("tables", []):
#                 st.dataframe(table)
#             for figure in content1.get(header, {}).get("figures", []):
#                 st.image(figure)
#         with col2:
#             st.subheader(header)
#             for para in content2.get(header, {}).get("paragraphs", []):
#                 st.write(para)
#             for table in content2.get(header, {}).get("tables", []):
#                 st.dataframe(table)
#             for figure in content2.get(header, {}).get("figures", []):
#                 st.image(figure)

# def main():
#     st.title("Word Document Side-By-Side Display with Tables and Figures")

#     uploaded_files = st.file_uploader("Choose two Word documents", accept_multiple_files=True, type='docx')

#     if len(uploaded_files) == 2:
#         content1 = load_document(uploaded_files[0])
#         content2 = load_document(uploaded_files[1])
#         display_side_by_side(content1, content2)

# if __name__ == "__main__":
#     main()

def display_multiple_documents(contents_by_header):
    number_of_documents = len(contents_by_header)
    cols = st.columns(number_of_documents)

    for header, content_dictionaries in contents_by_header.items():
        for i, content in enumerate(content_dictionaries):
            with cols[i]:
                st.subheader(header)
                for para in content.get("paragraphs", []):
                    st.write(para)
                for table in content.get("tables", []):
                    st.dataframe(table)
                for figure in content.get("figures", []):
                    st.image(figure, use_column_width=True)

def main():
    st.title("Word Document Multiple Comparison")

    uploaded_files = st.file_uploader(
        "Choose Word documents",
        accept_multiple_files=True,
        type='docx'
    )

    if uploaded_files and len(uploaded_files) > 1:
        contents = [load_document(uploaded_file) for uploaded_file in uploaded_files]

        # Find unique headers across documents
        all_headers = set()
        for content in contents:
            all_headers.update(content.keys())

        # Inject custom CSS with markdown to enable horizontal scrolling
        st.markdown("""
        <style>
            .st-eb {
                overflow-x: auto;
            }
            .wide-container {
                min-width: 3000px;
            }
        </style>
        """, unsafe_allow_html=True)

        # Create a container for the horizontally scrollable section
        with st.container():
            # Apply the custom class to this container
            st.markdown('<div class="st-eb">', unsafe_allow_html=True)
            st.markdown('<div class="wide-container">', unsafe_allow_html=True)

            for header in sorted(all_headers):
                cols = st.columns(len(uploaded_files))
                for i, content in enumerate(contents):
                    with cols[i]:
                        st.subheader(header)
                        for para in content.get(header, {}).get("paragraphs", []):
                            st.write(para)
                        for table in content.get(header, {}).get("tables", []):
                            st.dataframe(table)
                        for figure in content.get(header, {}).get("figures", []):
                            st.image(figure, use_column_width=True)

            st.markdown('</div>', unsafe_allow_html=True)  # end wide-container
            st.markdown('</div>', unsafe_allow_html=True)  # end overflow container

if __name__ == "__main__":
    main()